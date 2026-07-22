"""本地 AI 语音聊天:每个可见文本片段均先完成 IndexTTS2 合成."""
from __future__ import annotations

import base64
import configparser
import ipaddress
import io
import json
import logging
import math
import os
import queue
import re
import socket
import sys
import threading
import time
import uuid
from html import unescape
from pathlib import Path
from typing import Iterator
from urllib.parse import parse_qs, quote, quote_plus, unquote, urlencode, urljoin, urlparse

import httpx
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from langchain_core.messages import AIMessage, BaseMessage, HumanMessage, SystemMessage, ToolMessage
from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field


APP_DIR = Path(__file__).resolve().parent
STATIC_DIR = APP_DIR / "static"
MEDIA_DIR = APP_DIR / "runtime" / "audio"
MEDIA_DIR.mkdir(parents=True, exist_ok=True)
GENERATED_IMAGE_DIR = MEDIA_DIR.parent / "images"
GENERATED_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
LIVE2D_MODELS_DIR = APP_DIR / "live2dmodels"
CONVERSATIONS_PATH = MEDIA_DIR.parent / "conversations.json"
CONVERSATIONS_LOCK = threading.Lock()
DEFAULT_INDEXTTS_HOME = Path(r"D:\yzylauncher-win-Indextts20-260616\win-unpacked\python")
EMOTION_PROMPT_PATH = APP_DIR / "emotion_output_prompt.txt"
LIVE2D_CONTROL_PROMPT_PATH = APP_DIR / "live2d_output_prompt.txt"
EMOTION_DELIMITER = "&&"
LIVE2D_CONTROL_DELIMITER = "##"
DEFAULT_EMOTION_VECTOR = [0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.6]
# Merge consecutive very short streaming rows into a reasonably sized TTS job.
# Values count visible non-whitespace characters.
TTS_BATCH_TARGET_CHARS = 48
TTS_BATCH_MAX_CHARS = 64
TTS_SHORT_LINE_CHARS = 20
TTS_BATCH_MAX_WAIT_SECONDS = 0.55
ATTACHMENTS_LOCK = threading.Lock()
ATTACHMENTS: dict[str, dict] = {}
MAX_ATTACHMENT_BYTES = 10 * 1024 * 1024
MAX_ATTACHMENTS_PER_MESSAGE = 5
MAX_ATTACHMENT_TEXT_CHARS = 24_000
ATTACHMENT_TTL_SECONDS = 2 * 60 * 60
TEXT_FILE_EXTENSIONS = {
    ".txt", ".md", ".csv", ".json", ".log", ".ini", ".yaml", ".yml", ".xml",
    ".py", ".js", ".ts", ".html", ".css", ".sql", ".java", ".c", ".cpp", ".h",
    ".hpp", ".sh", ".ps1", ".bat", ".cmd",
}
IMAGE_FILE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
AUDIO_RETENTION_COUNT = 10
IMAGE_MODEL_MARKERS = (
    "gpt-image", "dall-e", "gemini", "qwen-image", "wanx", "flux", "recraft",
    "stable-diffusion", "sdxl", "z-image", "kolors", "cogview", "ideogram", "seedream",
)
IMAGE_MODEL_FALLBACK_ORDER = (
    "gpt-image-2","gpt-image-1.5", "gpt-image-1", "gemini-3.1-flash-image", "gemini-3-pro-image",
    "gemini-2.5-flash-image", "qwen-image-2.0-pro", "qwen-image-max", "flux-2-pro",
    "recraft-v4", "dall-e-3",
)
AUTO_IMAGE_MODEL_CACHE: dict[str, str] = {}
AUTO_IMAGE_MODEL_LOCK = threading.Lock()


def configure_console_logging() -> logging.Logger:
    """Send application and library diagnostics to the same CMD window as Uvicorn."""
    level_name = os.environ.get("AI_VOICE_CHAT_LOG_LEVEL", "DEBUG").upper()
    level = getattr(logging, level_name, logging.DEBUG)
    root = logging.getLogger()
    if not root.handlers:
        handler = logging.StreamHandler(sys.stdout)
        handler.setFormatter(logging.Formatter("%(asctime)s [%(levelname)s] %(name)s: %(message)s"))
        root.addHandler(handler)
    root.setLevel(level)
    for name in ("ai_voice_chat", "httpx", "httpcore"):
        logging.getLogger(name).setLevel(level)
    return logging.getLogger("ai_voice_chat")


LOGGER = configure_console_logging()


def normalize_provider(value: str) -> str:
    normalized = value.strip().lower().replace("-", "_").replace(" ", "_")
    aliases = {
        "openai": "openai_compatible",
        "openai_compatible": "openai_compatible",
        "openai_api": "openai_compatible",
        "lmstudio": "lm_studio",
        "lm_studio": "lm_studio",
    }
    return aliases.get(normalized, normalized)


def prune_generated_audio(directory: Path = MEDIA_DIR, keep: int = AUDIO_RETENTION_COUNT) -> int:
    """Retain only the newest generated WAV files; leave other media untouched."""
    files = sorted(
        (item for item in directory.glob("*.wav") if item.is_file()),
        key=lambda item: item.stat().st_mtime_ns,
        reverse=True,
    )
    removed = 0
    for item in files[max(0, keep):]:
        try:
            item.unlink()
            removed += 1
        except OSError as error:
            print(f"[TTS] unable to remove old audio {item.name}: {error}", flush=True)
    return removed


def load_prompt(path: Path, label: str) -> str:
    if not path.is_file():
        raise RuntimeError(f"找不到{label}:{path}")
    return path.read_text(encoding="utf-8").strip()


SYSTEM_PROMPT = load_prompt(EMOTION_PROMPT_PATH, "情绪输出约束文件")
LIVE2D_CONTROL_PROMPT = load_prompt(LIVE2D_CONTROL_PROMPT_PATH, "Live2D 输出约束文件")
TOOL_SYSTEM_PROMPT = """
你是一个会自主规划和执行的 AI Agent.先理解用户目标,再判断是否需要使用工具;每次工具返回后评估结果,必要时再选择下一步.
遇到需要最新、可核实或有来源的信息时先调用 web_search;搜索结果只有摘要,需要事实细节、原文内容或确认来源时必须再调用 browse_webpage 读取相关网页正文,不能把搜索摘要当作完整证据.
若网页工具报告网站拒绝自动读取,应改用搜索结果中的其他公开来源继续核实,不要直接结束任务或声称已读到该网页.
用户询问当前位置附近的天气、温度、环境或设施时,若本轮消息包含已授权定位信息,调用 get_local_environment;若没有定位信息,明确请用户点击“位置”按钮授权,不要猜测城市或位置.
用户明确要求创作、生成或绘制图片时调用 generate_image.不要为了普通闲聊、稳定常识或仅仅提及图片而调用工具,也不要在用户未提出图片需求时生成图片.
工具调用完成后,用工具结果回答用户;最终可见回复仍必须严格遵守每行 `文字&&[8个情感值]` 的格式.不要向用户输出工具调用 JSON.
绝不把 JSON、函数参数、tool_calls、function_call、事件名、URL 参数或工具原始返回直接写入可见回复；这些都属于内部信息。只提炼用户需要的结论、来源和清晰的自然语言说明。
""".strip()
MAX_AGENT_TOOL_ROUNDS = 10
MAX_AGENT_IMAGE_CALLS = 10
MAX_AGENT_SEARCH_CALLS = 10
MAX_AGENT_BROWSE_CALLS = 10
MAX_AGENT_LOCATION_CALLS = 10
AUTOMATION_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "搜索互联网的最新公开信息,并返回标题、摘要和来源链接.",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string", "description": "要搜索的准确关键词或问题"}},
                "required": ["query"],
                "additionalProperties": False,
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "browse_webpage",
            "description": "读取一个公开网页的标题与正文。应优先使用 web_search 返回的链接，并用它核实重要事实。",
            "parameters": {
                "type": "object",
                "properties": {"url": {"type": "string", "description": "要读取的 http 或 https 网页链接"}},
                "required": ["url"],
                "additionalProperties": False,
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_local_environment",
            "description": "基于用户本轮已授权的位置，查询当前天气、温度和附近公共设施。无授权位置时不可调用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "用户关注的附近设施或环境，例如咖啡馆、医院、公交、适合散步"},
                    "radius_m": {"type": "integer", "description": "查询半径，300 到 5000 米，默认 1500"},
                },
                "required": ["query"],
                "additionalProperties": False,
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "根据用户明确的画图、生成图片或设计图片请求创建一张图片.",
            "parameters": {
                "type": "object",
                "properties": {"prompt": {"type": "string", "description": "适合图片生成 API 的详细画面提示词"}},
                "required": ["prompt"],
                "additionalProperties": False,
            },
        },
    },
]


def load_settings() -> dict:
    """只读取同目录 config.txt,集中管理 API 和本地运行环境."""
    path = APP_DIR / "config.txt"
    if not path.is_file():
        raise RuntimeError(f"找不到配置文件:{path}")
    config = configparser.ConfigParser(interpolation=None)
    config.read(path, encoding="utf-8")

    def value(section: str, option: str, default: str | None = None) -> str:
        if config.has_option(section, option):
            return config.get(section, option).strip()
        if default is not None:
            return default
        raise RuntimeError(f"config.txt 缺少 [{section}] {option} 配置.")

    def bounded_int(section: str, option: str, default: int, minimum: int, maximum: int) -> int:
        try:
            number = int(value(section, option, str(default)))
        except ValueError as error:
            raise RuntimeError(f"config.txt 的 [{section}] {option} 必须是整数.") from error
        return max(minimum, min(maximum, number))

    base_url = value("chat", "base_url", "").rstrip("/")
    if base_url and re.search(r"/(?:models|chat/completions)$", base_url, flags=re.IGNORECASE):
        raise RuntimeError("[chat] base_url 必须是 API 根路径,例如 http://127.0.0.1:1234/v1,不能填写 /models 或 /chat/completions.")

    thinking_value = value("chat", "thinking", "").lower()
    if thinking_value and thinking_value not in {"true", "false", "1", "0", "yes", "no", "on", "off"}:
        raise RuntimeError("[chat] thinking must be true or false.")
    thinking = None if not thinking_value else thinking_value in {"true", "1", "yes", "on"}
    image_base_url = value("image", "base_url", "").rstrip("/")
    if image_base_url and re.search(r"/(?:models|images/generations)$", image_base_url, flags=re.IGNORECASE):
        raise RuntimeError("[image] base_url 必须是 API 根路径,例如 https://api.openai.com/v1.")
    image_configured = any(
        config.has_option("image", option) and config.get("image", option).strip()
        for option in ("api_key", "base_url", "model", "api_mode", "responses_model")
    )

    return {
        "index_tts_home": value("index_tts", "project_home"),
        "default_voice": value("index_tts", "default_voice", ""),
        "playback_prebuffer_segments": bounded_int("index_tts", "playback_prebuffer_segments", 2, 1, 4),
        "chat": {
            "provider": normalize_provider(value("chat", "provider", "openai_compatible")),
            "api_key": value("chat", "api_key", ""),
            "base_url": base_url,
            "model": value("chat", "model", ""),
            "thinking": thinking,
        },
        "image": {
            "api_key": value("image", "api_key", ""),
            "base_url": image_base_url,
            "model": value("image", "model", "auto") or "auto",
            "size": value("image", "size", "1024x1024") or "1024x1024",
            "quality": value("image", "quality", "high").lower() or "high",
            "api_mode": value("image", "api_mode", "auto").lower() or "auto",
            "responses_model": value("image", "responses_model", ""),
            "configured": image_configured,
        },
    }


SETTINGS = load_settings()


class WebChatSettings(BaseModel):
    provider: str = Field(default="openai_compatible", max_length=40)
    api_key: str = Field(default="", max_length=1000)
    base_url: str = Field(default="", max_length=1000)
    model: str = Field(default="", max_length=300)
    thinking: bool = False
    thinking_override: bool = False
    force_web_config: bool = False


class WebImageSettings(BaseModel):
    api_key: str = Field(default="", max_length=1000)
    base_url: str = Field(default="", max_length=1000)
    model: str = Field(default="auto", max_length=300)
    api_mode: str = Field(default="auto", max_length=20)
    responses_model: str = Field(default="", max_length=300)
    force_web_config: bool = False


class LocationContext(BaseModel):
    """Ephemeral browser location supplied only after the user grants permission."""
    latitude: float = Field(ge=-90, le=90)
    longitude: float = Field(ge=-180, le=180)
    accuracy: float = Field(default=0, ge=0, le=100_000)


class AttachmentReference(BaseModel):
    id: str = Field(min_length=32, max_length=64)


class ChatRequest(BaseModel):
    message: str = Field(min_length=1, max_length=4000)
    history: list[dict[str, str]] = Field(default_factory=list, max_length=12)
    attachments: list[AttachmentReference] = Field(default_factory=list, max_length=MAX_ATTACHMENTS_PER_MESSAGE)
    web_search: bool = False
    location: LocationContext | None = None
    voice: str = Field(min_length=1, max_length=160)
    speaking_speed: float = Field(default=1.0, ge=0.75, le=1.35)
    agent_enabled: bool = True
    live2d_model_id: str = Field(default="", max_length=1000)
    web_chat: WebChatSettings = Field(default_factory=WebChatSettings)
    web_image: WebImageSettings = Field(default_factory=WebImageSettings)


class ImageRequest(BaseModel):
    prompt: str = Field(min_length=1, max_length=4000)
    web_chat: WebChatSettings = Field(default_factory=WebChatSettings)
    web_image: WebImageSettings = Field(default_factory=WebImageSettings)


def resolve_chat_config_from_web(web: WebChatSettings) -> dict:
    configured = SETTINGS["chat"]
    configured_provider = str(configured.get("provider", "openai_compatible")).strip().lower()
    configured_complete = bool(
        configured.get("base_url") and configured.get("model")
        and (configured.get("api_key") or configured_provider == "lm_studio")
    )

    def select(name: str) -> str:
        configured_value = str(configured.get(name, "")).strip()
        web_value = str(getattr(web, name)).strip()
        return web_value if web.force_web_config else configured_value or web_value

    provider = normalize_provider(str(web.provider).strip() if web.force_web_config or not configured_complete else configured_provider) or "openai_compatible"
    if provider not in {"openai_compatible", "lm_studio"}:
        raise RuntimeError("Unsupported provider. Choose OpenAI compatible or LM Studio.")
    base_url = select("base_url").rstrip("/")
    model = select("model")
    api_key = select("api_key")
    thinking = web.thinking if web.thinking_override or web.force_web_config or not configured_complete or configured.get("thinking") is None else bool(configured["thinking"])

    if not base_url or not model:
        raise RuntimeError("请在 config.txt 或网页模型设置中填写 base_url 和模型名称.")
    if re.search(r"/(?:models|chat/completions)$", base_url, flags=re.IGNORECASE):
        raise RuntimeError("base_url 必须填写 API 根路径,不能填写 /models 或 /chat/completions.")
    if provider != "lm_studio" and not api_key:
        raise RuntimeError("请在 config.txt 或网页模型设置中填写 API Key.")
    return {"provider": provider, "base_url": base_url, "model": model, "api_key": api_key, "thinking": thinking}


def resolve_chat_config(request: ChatRequest) -> dict:
    return resolve_chat_config_from_web(request.web_chat)


def image_api_headers(api_key: str) -> dict[str, str]:
    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    return headers


def image_api_uses_env_proxy(base_url: str) -> bool:
    """Keep localhost and LAN OpenAI-compatible services away from broken SOCKS settings."""
    host = (urlparse(base_url).hostname or "").strip().lower()
    if host in {"", "localhost"}:
        return False
    try:
        address = ipaddress.ip_address(host)
    except ValueError:
        return True
    return not (address.is_private or address.is_loopback or address.is_link_local)


def image_model_version(model_id: str) -> tuple[int, ...]:
    return tuple(int(value) for value in re.findall(r"\d+", model_id))


def is_image_model(model_id: str) -> bool:
    normalized = model_id.strip().lower()
    return bool(normalized) and any(marker in normalized for marker in IMAGE_MODEL_MARKERS)


def select_latest_image_model(models: list[dict]) -> str:
    """Pick a non-preview image model from the endpoint's own model list."""
    candidates = [item for item in models if is_image_model(str(item.get("id", "")))]
    if not candidates:
        return ""

    preferred = {model: len(IMAGE_MODEL_FALLBACK_ORDER) - index for index, model in enumerate(IMAGE_MODEL_FALLBACK_ORDER)}

    def ranking(item: dict) -> tuple[int, int, int, tuple[int, ...]]:
        model_id = str(item.get("id", "")).strip().lower()
        stable = 0 if "preview" in model_id or "deprecated" in model_id else 1
        try:
            created = int(item.get("created", 0) or 0)
        except (TypeError, ValueError):
            created = 0
        return stable, created, preferred.get(model_id, 0), image_model_version(model_id)

    return str(max(candidates, key=ranking).get("id", "")).strip()


def discover_image_model(base_url: str, api_key: str) -> str:
    cache_key = f"{base_url}\x00{api_key}"
    with AUTO_IMAGE_MODEL_LOCK:
        if cached := AUTO_IMAGE_MODEL_CACHE.get(cache_key):
            return cached

    endpoint = f"{base_url}/models"
    try:
        with httpx.Client(timeout=20, trust_env=image_api_uses_env_proxy(base_url)) as client:
            response = client.get(endpoint, headers=image_api_headers(api_key))
            response.raise_for_status()
        payload = response.json()
    except (httpx.HTTPError, ValueError):
        return ""

    raw_models = payload.get("data", []) if isinstance(payload, dict) else []
    models = [item if isinstance(item, dict) else {"id": item} for item in raw_models if isinstance(item, (dict, str))]
    selected = select_latest_image_model(models)
    if selected:
        with AUTO_IMAGE_MODEL_LOCK:
            AUTO_IMAGE_MODEL_CACHE[cache_key] = selected
    return selected


def resolve_image_config(web_image: WebImageSettings | None = None, chat_config: dict | None = None) -> dict:
    configured = SETTINGS["image"]
    web = web_image or WebImageSettings()
    chat_connection = chat_config or SETTINGS["chat"]
    configured_is_present = bool(configured.get("configured"))

    def select(name: str) -> str:
        configured_value = str(configured.get(name, "")).strip()
        web_value = str(getattr(web, name)).strip()
        return web_value if web.force_web_config or not configured_is_present else configured_value or web_value

    # Image endpoints require a dedicated image model. Never fall back to the
    # chat model: it produces a confusing HTTP 400 from providers such as OpenAI.
    base_url = select("base_url").rstrip("/") or str(chat_connection.get("base_url", "")).strip().rstrip("/")
    configured_model = select("model") or "auto"
    api_mode = select("api_mode").lower() or "auto"
    responses_model = select("responses_model") or str(chat_connection.get("model", "")).strip()
    api_key = select("api_key") or str(chat_connection.get("api_key", "")).strip()
    if not base_url:
        raise RuntimeError("未配置画图 API 的 base_url.请填写 [chat] 或 [image] 的 base_url.")
    if re.search(r"/(?:models|images/generations)$", base_url, flags=re.IGNORECASE):
        raise RuntimeError("[image] base_url 必须填写 API 根路径,不能填写 /models 或 /images/generations.")
    if api_mode not in {"auto", "images", "responses"}:
        raise RuntimeError("[image] api_mode 只能是 auto、images 或 responses.")
    quality = str(configured.get("quality", "high")).strip().lower() or "high"
    if quality not in {"auto", "low", "medium", "high"}:
        raise RuntimeError("[image] quality 只能是 auto、low、medium 或 high.")
    if api_mode == "responses" and not responses_model:
        raise RuntimeError("Responses 图片工具需要填写 Responses 模型，或先填写聊天模型名称。")
    if configured_model.lower() in {"auto", "latest"}:
        model = discover_image_model(base_url, api_key)
        if not model:
            raise RuntimeError(
                "未能从当前接口的 /models 自动识别图片模型.请在 config.txt 的 [image] 中填写 model,"
                "例如 OpenAI 使用 gpt-image-1.5,Gemini OpenAI 兼容接口使用 gemini-3.1-flash-image."
            )
    else:
        model = configured_model
    return {
        "base_url": base_url,
        "model": model,
        "api_key": api_key,
        "size": str(configured.get("size", "1024x1024")).strip() or "1024x1024",
        "quality": quality,
        "api_mode": api_mode,
        "responses_model": responses_model,
    }


def image_format(data: bytes) -> tuple[str, str]:
    if data.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png", "image/png"
    if data.startswith(b"\xff\xd8\xff"):
        return ".jpg", "image/jpeg"
    if data.startswith((b"GIF87a", b"GIF89a")):
        return ".gif", "image/gif"
    if data.startswith(b"RIFF") and data[8:12] == b"WEBP":
        return ".webp", "image/webp"
    return ".png", "image/png"


def save_generated_image(data: bytes) -> str:
    suffix, _ = image_format(data)
    filename = f"{int(time.time() * 1000)}-{uuid.uuid4().hex}{suffix}"
    (GENERATED_IMAGE_DIR / filename).write_bytes(data)
    return f"/generated-images/{filename}"


def image_from_responses_payload(result: dict) -> dict[str, str]:
    """Read the image_generation tool output from the OpenAI Responses API."""
    output = result.get("output", []) if isinstance(result, dict) else []
    call = next(
        (
            item for item in output
            if isinstance(item, dict) and item.get("type") == "image_generation_call" and isinstance(item.get("result"), str)
        ),
        None,
    )
    if not call:
        raise RuntimeError("Responses API 未返回 image_generation 工具结果。")
    try:
        data = base64.b64decode(call["result"], validate=True)
    except ValueError as error:
        raise RuntimeError("Responses API 返回了无法解码的图片数据。") from error
    return {"url": save_generated_image(data), "revised_prompt": ""}


def image_from_responses_sse_lines(lines: Iterator[str]) -> dict[str, str]:
    """Extract the final image from streaming Responses API events."""
    completed_response: dict | None = None
    partial_image_b64 = ""
    completed_item: dict | None = None
    event_types: list[str] = []
    upstream_error = ""
    for raw_line in lines:
        line = raw_line.strip()
        if not line.startswith("data:"):
            continue
        data = line[5:].strip()
        if not data or data == "[DONE]":
            continue
        try:
            event = json.loads(data)
        except json.JSONDecodeError:
            continue
        if not isinstance(event, dict):
            continue
        event_type = str(event.get("type", "")).strip()
        if event_type and event_type not in event_types:
            event_types.append(event_type)
        if event_type == "response.image_generation_call.partial_image":
            candidate = event.get("partial_image_b64")
            if isinstance(candidate, str) and candidate:
                partial_image_b64 = candidate
        elif event_type == "response.completed" and isinstance(event.get("response"), dict):
            completed_response = event["response"]
        elif event_type == "response.output_item.done" and isinstance(event.get("item"), dict):
            completed_item = event["item"]
        elif event_type in {"response.failed", "error"}:
            error = event.get("error")
            if isinstance(error, dict) and isinstance(error.get("message"), str):
                upstream_error = error["message"]

    if completed_item:
        try:
            return image_from_responses_payload({"output": [completed_item]})
        except RuntimeError:
            pass

    if completed_response:
        try:
            return image_from_responses_payload(completed_response)
        except RuntimeError:
            pass
    if not partial_image_b64:
        diagnostic = ", ".join(event_types[-8:]) or "无可解析事件"
        print(f"[Image] Responses stream had no image result; events={diagnostic}", flush=True)
        if upstream_error:
            raise RuntimeError(f"Responses 图片工具被上游拒绝：{upstream_error}")
        raise RuntimeError(f"Responses API 未返回 image_generation 工具结果（事件：{diagnostic}）。")
    try:
        data = base64.b64decode(partial_image_b64, validate=True)
    except ValueError as error:
        raise RuntimeError("Responses API 返回了无法解码的图片数据。") from error
    return {"url": save_generated_image(data), "revised_prompt": ""}


def response_image_input(prompt: str, reference_images: list[dict] | None = None) -> str | list[dict]:
    """Include source images in the generation request so edits preserve them."""
    images = (reference_images or [])[:4]
    if not images:
        return prompt
    content: list[dict] = [{
        "type": "input_text",
        "text": (
            "Edit the supplied source image. Preserve the subject identity, face, pose, framing, "
            "composition, background, lighting, and style. Change only what the instruction requests; "
            "do not redesign unrelated details.\n\nInstruction:\n" + prompt
        ),
    }]
    for image in images:
        content_type = str(image.get("content_type", "image/png"))
        encoded = image.get("data")
        if isinstance(encoded, str) and encoded:
            content.append({"type": "input_image", "image_url": f"data:{content_type};base64,{encoded}"})
    return [{"role": "user", "content": content}]


def generate_image_via_responses(
    prompt: str,
    config: dict,
    reference_images: list[dict] | None = None,
    include_partial_images: bool = True,
) -> dict[str, str]:
    """Use streaming Responses image tools, including Sub2API's gpt-image-2 path."""
    endpoint = f"{config['base_url']}/responses"
    image_tool = {
        "type": "image_generation",
        "quality": config["quality"],
        "size": config["size"],
    }
    if include_partial_images:
        image_tool["partial_images"] = 1
    payload = {
        "model": config["responses_model"],
        "input": response_image_input(prompt, reference_images),
        # A partial image helps gateways that omit `result` in response.completed.
        "tools": [image_tool],
        "tool_choice": "required",
        "stream": True,
    }
    try:
        with httpx.Client(timeout=300, trust_env=image_api_uses_env_proxy(config["base_url"])) as client:
            with client.stream("POST", endpoint, headers=image_api_headers(config["api_key"]), json=payload) as response:
                response.raise_for_status()
                return image_from_responses_sse_lines(response.iter_lines())
    except httpx.HTTPStatusError as error:
        detail = error.response.text[:500].strip()
        raise RuntimeError(f"Responses 图片工具调用失败(HTTP {error.response.status_code})：{detail}") from error
    except httpx.HTTPError as error:
        raise RuntimeError(f"Responses 图片工具连接失败：{error}") from error
    except RuntimeError as error:
        # Some Sub2API upstreams abort while forwarding partial-image SSE
        # chunks. Retry once with the same streaming protocol but only the
        # final image event, avoiding a second expensive retry loop.
        if include_partial_images and "stream_read_error" in str(error).lower():
            print("[Image] upstream stream_read_error; retrying Responses without partial images.", flush=True)
            return generate_image_via_responses(prompt, config, reference_images, include_partial_images=False)
        raise
def generate_image(prompt: str, config: dict, reference_images: list[dict] | None = None) -> dict[str, str]:
    if reference_images:
        print(
            f"[Image] editing {len(reference_images)} source image(s) via Responses "
            f"model={config['responses_model']} quality={config['quality']}",
            flush=True,
        )
        return generate_image_via_responses(prompt, config, reference_images)
    if config.get("api_mode") == "responses":
        print(f"[Image] using streaming Responses tool model={config['responses_model']} quality={config['quality']}", flush=True)
        return generate_image_via_responses(prompt, config)
    endpoint = f"{config['base_url']}/images/generations"
    headers = image_api_headers(config["api_key"])
    payload = {"model": config["model"], "prompt": prompt, "n": 1, "size": config["size"], "response_format": "b64_json"}
    try:
        with httpx.Client(timeout=300, trust_env=image_api_uses_env_proxy(config["base_url"])) as client:
            response = client.post(endpoint, headers=headers, json=payload)
            # Some OpenAI-compatible providers only implement URL responses.
            if response.status_code == 400 and "response_format" in response.text.lower():
                payload.pop("response_format", None)
                response = client.post(endpoint, headers=headers, json=payload)
            response.raise_for_status()
    except httpx.HTTPStatusError as error:
        detail = error.response.text[:500].strip()
        detail_lower = detail.lower()
        if error.response.status_code == 400 and "tool choice" in detail_lower and "image_generation" in detail_lower:
            print("[Image] switching to Responses image_generation tool protocol.", flush=True)
            return generate_image_via_responses(prompt, config, reference_images)
        if (
            config.get("api_mode") == "auto"
            and error.response.status_code == 502
            and config["model"].strip().lower() in {"gpt-image-2", "codex-gpt-image-2"}
            and "upstream request failed" in detail_lower
        ):
            print("[Image] gpt-image-2 upstream 502; retrying via streaming Responses image tool.", flush=True)
            return generate_image_via_responses(prompt, config, reference_images)
        if error.response.status_code == 400 and "requires an image model" in detail.lower():
            raise RuntimeError(
                f"画图模型「{config['model']}」不是图片模型.请将 config.txt 的 [image] model 设为 auto,"
                "或填写该服务商提供的专用图片模型."
            ) from error
        raise RuntimeError(f"画图 API 调用失败(HTTP {error.response.status_code}):{detail}") from error
    except httpx.HTTPError as error:
        raise RuntimeError(f"画图 API 连接失败:{error}") from error

    try:
        result = response.json()
        image = result["data"][0]
    except (KeyError, IndexError, TypeError, ValueError) as error:
        raise RuntimeError("画图 API 返回格式不正确,未找到图片数据.") from error

    image_url = image.get("url") if isinstance(image, dict) else None
    if isinstance(image, dict) and isinstance(image.get("b64_json"), str):
        try:
            data = base64.b64decode(image["b64_json"], validate=True)
        except ValueError as error:
            raise RuntimeError("画图 API 返回了无法解码的图片数据.") from error
        image_url = save_generated_image(data)
    if not isinstance(image_url, str) or not image_url:
        raise RuntimeError("画图 API 未返回图片地址.")
    revised_prompt = image.get("revised_prompt", "") if isinstance(image, dict) else ""
    return {"url": image_url, "revised_prompt": revised_prompt if isinstance(revised_prompt, str) else ""}


class ConversationMessage(BaseModel):
    role: str = Field(min_length=1, max_length=16)
    content: str = Field(min_length=1, max_length=10000)
    model_content: str | None = Field(default=None, max_length=12000)


class ConversationSaveRequest(BaseModel):
    id: str = Field(min_length=1, max_length=80)
    title: str = Field(min_length=1, max_length=120)
    messages: list[ConversationMessage] = Field(min_length=1, max_length=120)


class IndexTTSService:
    """按需加载一个 IndexTTS2 实例,避免同时生成时争用显卡."""

    def __init__(self, project_home: Path):
        self.project_home = project_home
        self.voices_dir = project_home / "voices"
        self.model_dir = project_home / "checkpoints"
        self._model = None
        self._load_lock = threading.Lock()
        self._infer_lock = threading.Lock()
        self._model_state = "not_loaded"
        self._model_error = ""

    def voices(self) -> list[str]:
        if not self.voices_dir.is_dir():
            return []
        extensions = {".wav", ".mp3", ".flac", ".ogg", ".m4a", ".aac"}
        return sorted(item.name for item in self.voices_dir.iterdir() if item.suffix.lower() in extensions)

    def _get_model(self):
        if self._model is None:
            with self._load_lock:
                if self._model is None:
                    self._model_state = "loading"
                    try:
                        if not self.model_dir.joinpath("config.yaml").is_file():
                            raise RuntimeError(f"找不到 IndexTTS2 模型目录:{self.model_dir}")
                        if str(self.project_home) not in sys.path:
                            sys.path.insert(0, str(self.project_home))
                        import torch
                        from indextts.infer_v2 import IndexTTS2

                        self._model = IndexTTS2(
                            cfg_path=str(self.model_dir / "config.yaml"),
                            model_dir=str(self.model_dir),
                            use_fp16=torch.cuda.is_available(),
                            use_cuda_kernel=False,
                            use_deepspeed=False,
                            low_vram=False,
                        )
                    except Exception as error:
                        self._model_state = "error"
                        self._model_error = str(error)
                        raise
                    self._model_state = "ready"
        return self._model

    def model_status(self) -> str:
        return self._model_state

    def synthesize(self, text: str, voice_name: str, emotion_vector: list[float], speaking_speed: float) -> str:
        # 文件名只允许来自 voices 目录的白名单,杜绝路径穿越.
        if voice_name not in self.voices():
            raise ValueError("所选音色不存在,请刷新页面后重新选择.")
        target = MEDIA_DIR / f"{int(time.time() * 1000)}-{uuid.uuid4().hex}.wav"
        with self._infer_lock:
            model = self._get_model()
            normalized_vector = model.normalize_emo_vec(emotion_vector, apply_bias=True)
            model.infer(
                spk_audio_prompt=str(self.voices_dir / voice_name),
                text=text,
                output_path=str(target),
                emo_vector=normalized_vector,
                use_random=False,
                verbose=False,
                # External streaming keeps each job short. 90 tokens is within IndexTTS2's
                # recommended 80-200 range and avoids an unnecessary second split here.
                max_text_tokens_per_segment=90,
                # Punctuation already carries a natural pause. Do not add artificial silence
                # between the small streaming jobs.
                interval_silence=0,
                speaking_speed=speaking_speed,
            )
        if not target.is_file():
            raise RuntimeError("IndexTTS2 没有生成音频文件.")
        removed = prune_generated_audio()
        if removed:
            print(f"[TTS] cleaned {removed} old WAV file(s); keeping newest {AUDIO_RETENTION_COUNT}.", flush=True)
        return target.name


TTS = IndexTTSService(Path(SETTINGS["index_tts_home"]))
app = FastAPI(title="AI Voice Chat")


def is_within_live2d_models(path: Path) -> bool:
    """Keep the local model endpoint limited to the configured model folder."""
    try:
        path.relative_to(LIVE2D_MODELS_DIR.resolve())
        return True
    except ValueError:
        return False


def completed_live2d_model_definition(model_path: Path) -> dict:
    """Register side-car expressions and motions omitted by some VTube exports."""
    definition = json.loads(model_path.read_text(encoding="utf-8"))
    references = definition.setdefault("FileReferences", {})
    model_dir = model_path.parent

    expressions_dir = model_dir / "expressions"
    if expressions_dir.is_dir() and not references.get("Expressions"):
        expressions = []
        for expression_path in sorted(expressions_dir.rglob("*.exp3.json")):
            expressions.append({
                "Name": expression_path.name.removesuffix(".exp3.json"),
                "File": expression_path.relative_to(model_dir).as_posix(),
            })
        if expressions:
            references["Expressions"] = expressions

    motions_dir = model_dir / "motions"
    if motions_dir.is_dir() and not references.get("Motions"):
        idle_motions = []
        manual_motions = []
        for motion_path in sorted(motions_dir.rglob("*.motion3.json")):
            motion = {"File": motion_path.relative_to(model_dir).as_posix()}
            # Only the intentionally named idle animation may loop by itself.
            # VTube exports often include pose-switch motions such as 变芒/变荒;
            # registering them as Idle makes the renderer trigger them randomly.
            if "待机" in motion_path.stem:
                idle_motions.append(motion)
            else:
                manual_motions.append(motion)
        motion_groups = {}
        if idle_motions:
            motion_groups["Idle"] = idle_motions
        if manual_motions:
            motion_groups["Manual"] = manual_motions
        if motion_groups:
            references["Motions"] = motion_groups
    return definition


def local_live2d_models() -> list[dict[str, object]]:
    if not LIVE2D_MODELS_DIR.is_dir():
        return []
    models = []
    for model_path in sorted(LIVE2D_MODELS_DIR.rglob("*.model3.json")):
        relative_path = model_path.relative_to(LIVE2D_MODELS_DIR)
        try:
            definition = completed_live2d_model_definition(model_path)
        except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
            LOGGER.warning("[Live2D] skipped invalid model %s: %s", relative_path, exc)
            continue
        expressions = [item.get("Name") for item in definition.get("FileReferences", {}).get("Expressions", [])]
        motion_groups = {}
        for group_name, group_motions in definition.get("FileReferences", {}).get("Motions", {}).items():
            names = []
            for motion in group_motions:
                if not isinstance(motion, dict) or not isinstance(motion.get("File"), str):
                    continue
                names.append(Path(motion["File"]).name.removesuffix(".motion3.json"))
            if names:
                motion_groups[group_name] = names
        models.append({
            "id": relative_path.as_posix(),
            "name": model_path.parent.name,
            "url": "/live2dmodels/" + quote(relative_path.as_posix()),
            "expressions": [name for name in expressions if isinstance(name, str)],
            "motion_groups": motion_groups,
        })
    return models


def selected_live2d_model(model_id: str) -> dict[str, object] | None:
    models = local_live2d_models()
    if not models:
        return None
    normalized_id = model_id.strip()
    if normalized_id:
        return next((model for model in models if model["id"] == normalized_id), models[0])
    return models[0]


def live2d_control_catalog(model_id: str) -> tuple[list[str], list[str]]:
    model = selected_live2d_model(model_id)
    if not model:
        return [], []
    expressions = sorted({str(name) for name in model.get("expressions", []) if isinstance(name, str)})
    motion_groups = model.get("motion_groups", {})
    manual_motions = motion_groups.get("Manual", []) if isinstance(motion_groups, dict) else []
    motions = sorted({str(name) for name in manual_motions if isinstance(name, str)})
    return expressions, motions


def live2d_control_system_prompt(model_id: str) -> str:
    expressions, motions = live2d_control_catalog(model_id)
    if not expressions and not motions:
        return ""
    # The selected model can contain many cosmetic assets. A bounded catalog
    # keeps the prompt useful while still exposing the first-class controls.
    catalog = {
        "expressions": expressions[:80],
        "manual_motions": motions[:80],
    }
    return (
        f"{LIVE2D_CONTROL_PROMPT}\n\n"
        "Live2D 当前模型可用控制项（仅可使用下列精确名称）：\n"
        f"{json.dumps(catalog, ensure_ascii=False, separators=(',', ':'))}"
    )


def normalize_live2d_control(control: object, model_id: str) -> dict[str, str | None] | None:
    """Keep LLM avatar commands inside the currently loaded model catalog."""
    if not isinstance(control, dict):
        return None
    expressions, motions = live2d_control_catalog(model_id)
    expression = control.get("expression")
    motion = control.get("motion")
    normalized_expression = expression if isinstance(expression, str) and expression in expressions else None
    normalized_motion = motion if isinstance(motion, str) and motion in motions else None
    return {"expression": normalized_expression, "motion": normalized_motion}


@app.get("/live2dmodels/{asset_path:path}", include_in_schema=False)
def serve_local_live2d_asset(asset_path: str):
    """Serve local Cubism resources, augmenting the model descriptor when needed."""
    if not LIVE2D_MODELS_DIR.is_dir():
        raise HTTPException(status_code=404, detail="live2dmodels folder was not found.")
    requested_path = (LIVE2D_MODELS_DIR / asset_path).resolve()
    if not is_within_live2d_models(requested_path) or not requested_path.is_file():
        raise HTTPException(status_code=404, detail="Live2D asset was not found.")
    if requested_path.name.endswith(".model3.json"):
        try:
            return Response(
                content=json.dumps(completed_live2d_model_definition(requested_path), ensure_ascii=False),
                media_type="application/json",
                headers={"Cache-Control": "no-store, max-age=0"},
            )
        except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
            LOGGER.warning("[Live2D] failed to read %s: %s", requested_path.name, exc)
            raise HTTPException(status_code=500, detail="Live2D model descriptor could not be read.") from exc
    return FileResponse(requested_path, headers={"Cache-Control": "no-store, max-age=0"})


app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
app.mount("/media", StaticFiles(directory=MEDIA_DIR), name="media")
app.mount("/generated-images", StaticFiles(directory=GENERATED_IMAGE_DIR), name="generated-images")


@app.middleware("http")
async def log_http_request(request, call_next):
    """Log request lifecycle without exposing message bodies or credentials."""
    started_at = time.perf_counter()
    try:
        response = await call_next(request)
    except Exception:
        LOGGER.exception("[HTTP] %s %s failed", request.method, request.url.path)
        raise
    if request.url.path == "/" or request.url.path.startswith("/static/") or request.url.path.startswith("/live2dmodels/"):
        # The app is launched locally and updated in place.  Do not let an old
        # cached index.html or app.js hide a newly added front-end feature.
        response.headers["Cache-Control"] = "no-store, max-age=0"
    elapsed_ms = (time.perf_counter() - started_at) * 1000
    LOGGER.info("[HTTP] %s %s -> %s (%.0f ms)", request.method, request.url.path, response.status_code, elapsed_ms)
    return response


def sse(event: str, data: dict) -> bytes:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n".encode("utf-8")


def conversation_summary(record: dict) -> dict:
    return {
        "id": record["id"],
        "title": record["title"],
        "updated_at": record["updated_at"],
    }


def load_saved_conversations() -> list[dict]:
    if not CONVERSATIONS_PATH.is_file():
        return []
    try:
        data = json.loads(CONVERSATIONS_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as error:
        print(f"[History] unable to read saved conversations: {error}", flush=True)
        return []
    if not isinstance(data, list):
        return []
    return [record for record in data if isinstance(record, dict) and isinstance(record.get("id"), str)]


def write_saved_conversations(records: list[dict]) -> None:
    temporary_path = CONVERSATIONS_PATH.with_suffix(".tmp")
    temporary_path.write_text(json.dumps(records, ensure_ascii=False, separators=(",", ":")), encoding="utf-8")
    temporary_path.replace(CONVERSATIONS_PATH)


def validate_conversation_id(conversation_id: str) -> str:
    if not re.fullmatch(r"[A-Za-z0-9_-]{1,80}", conversation_id):
        raise HTTPException(status_code=400, detail="Invalid conversation id.")
    return conversation_id


def clean_expired_attachments() -> None:
    deadline = time.time() - ATTACHMENT_TTL_SECONDS
    with ATTACHMENTS_LOCK:
        expired = [attachment_id for attachment_id, item in ATTACHMENTS.items() if item["created_at"] < deadline]
        for attachment_id in expired:
            ATTACHMENTS.pop(attachment_id, None)


def attachment_kind(filename: str, content_type: str | None) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in IMAGE_FILE_EXTENSIONS and (not content_type or content_type.startswith("image/")):
        return "image"
    if suffix in TEXT_FILE_EXTENSIONS or (content_type or "").startswith("text/"):
        return "text"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    raise HTTPException(
        status_code=400,
        detail="暂支持图片(PNG/JPG/WEBP/GIF)、文本/代码、PDF 和 DOCX 文件.",
    )


def decode_text_file(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="该文本文件编码无法识别,请保存为 UTF-8 后重新上传.")


def extract_attachment_text(kind: str, data: bytes) -> str:
    if kind == "text":
        text = decode_text_file(data)
    elif kind == "pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            text = "\n".join((page.extract_text() or "") for page in reader.pages[:30])
        except Exception as error:
            raise HTTPException(status_code=400, detail=f"无法读取 PDF 内容:{error}") from error
    elif kind == "docx":
        try:
            from docx import Document

            document = Document(io.BytesIO(data))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
            text = "\n".join(paragraphs + table_cells)
        except Exception as error:
            raise HTTPException(status_code=400, detail=f"无法读取 Word 文档内容:{error}") from error
    else:
        return ""
    text = text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="文件中没有可供模型读取的文本内容.")
    return text[:MAX_ATTACHMENT_TEXT_CHARS]


def resolve_attachments(references: list[AttachmentReference]) -> list[dict]:
    clean_expired_attachments()
    if not references:
        return []
    attachment_ids = [reference.id for reference in references]
    if len(set(attachment_ids)) != len(attachment_ids):
        raise RuntimeError("同一附件不能重复发送.")
    with ATTACHMENTS_LOCK:
        resolved = [ATTACHMENTS.get(attachment_id) for attachment_id in attachment_ids]
    if any(item is None for item in resolved):
        raise RuntimeError("有附件已失效,请重新上传后再发送.")
    return [item for item in resolved if item is not None]


def strip_html(value: str) -> str:
    value = re.sub(r"<script[\s\S]*?</script>|<style[\s\S]*?</style>", " ", value, flags=re.IGNORECASE)
    value = re.sub(r"<[^>]+>", " ", value)
    return re.sub(r"\s+", " ", unescape(value)).strip()


def unwrap_search_url(href: str) -> str:
    href = unescape(href)
    if href.startswith("//"):
        href = f"https:{href}"
    parsed = urlparse(href)
    if parsed.netloc.endswith("duckduckgo.com"):
        forwarded = parse_qs(parsed.query).get("uddg")
        if forwarded:
            return forwarded[0]
    return href


def web_search(query: str) -> list[dict[str, str]]:
    """Get a compact, citation-friendly result list without exposing a search API key."""
    if not query.strip():
        return []
    url = f"https://html.duckduckgo.com/html/?q={quote_plus(query[:1000])}"
    headers = {"User-Agent": "Mozilla/5.0 (AI-Voice-Chat local demo)"}
    try:
        with httpx.Client(timeout=15, follow_redirects=True, headers=headers) as client:
            response = client.get(url)
            response.raise_for_status()
    except httpx.HTTPError as error:
        raise RuntimeError(f"联网搜索连接失败:{error}") from error

    links: list[tuple[str, str]] = []
    for anchor in re.findall(r"<a\b[^>]*>[\s\S]*?</a>", response.text, flags=re.IGNORECASE):
        if "result__a" not in anchor:
            continue
        match = re.search(r'''href=["']([^"']+)["']''', anchor, flags=re.IGNORECASE)
        if not match:
            continue
        title = strip_html(anchor)
        href = unwrap_search_url(match.group(1))
        if title and href.startswith(("http://", "https://")):
            links.append((title, href))

    snippets = [strip_html(item) for item in re.findall(
        r'''<[^>]*class=["'][^"']*result__snippet[^"']*["'][^>]*>([\s\S]*?)</(?:a|div)>''',
        response.text,
        flags=re.IGNORECASE,
    )]
    results: list[dict[str, str]] = []
    seen_urls: set[str] = set()
    for index, (title, href) in enumerate(links):
        if href in seen_urls:
            continue
        seen_urls.add(href)
        results.append({"title": title[:180], "url": href, "snippet": (snippets[index] if index < len(snippets) else "")[:500]})
        if len(results) == 5:
            break
    return results


def validate_public_web_url(url: str) -> str:
    """Allow public web browsing while blocking server-side access to local networks."""
    parsed = urlparse(url.strip())
    if parsed.scheme not in {"http", "https"} or not parsed.hostname or parsed.username or parsed.password:
        raise RuntimeError("只能浏览公开的 http 或 https 网页链接.")
    hostname = parsed.hostname.strip().lower()
    if hostname in {"localhost", "localhost.localdomain"}:
        raise RuntimeError("不能浏览本机或内网地址.")
    try:
        addresses = {item[4][0] for item in socket.getaddrinfo(hostname, parsed.port or 443, type=socket.SOCK_STREAM)}
    except OSError as error:
        raise RuntimeError(f"网页域名无法解析:{hostname}") from error
    try:
        if not addresses or any(not ipaddress.ip_address(address).is_global for address in addresses):
            raise RuntimeError("不能浏览本机或内网地址.")
    except ValueError as error:
        raise RuntimeError("网页地址无效.") from error
    return parsed.geturl()


def browse_mediawiki_api(url: str) -> dict[str, str] | None:
    """Use a site's documented public MediaWiki API when its HTML page blocks readers."""
    parsed = urlparse(url)
    page = unquote(parsed.path).strip("/")
    if not page or page.endswith("api.php"):
        return None
    api_url = f"{parsed.scheme}://{parsed.netloc}/api.php?" + urlencode({
        "action": "parse", "page": page, "prop": "text", "format": "json", "formatversion": "2",
    })
    try:
        api_url = validate_public_web_url(api_url)
        with httpx.Client(timeout=20, headers={"User-Agent": "AI-Voice-Chat local demo", "Accept": "application/json"}) as client:
            response = client.get(api_url)
            response.raise_for_status()
        payload = response.json()
        parsed_page = payload.get("parse", {}) if isinstance(payload, dict) else {}
        html = parsed_page.get("text", "") if isinstance(parsed_page, dict) else ""
        if not isinstance(html, str):
            return None
        text = strip_html(html)[:12_000]
        if not text:
            return None
        title = str(parsed_page.get("title", ""))[:240] if isinstance(parsed_page, dict) else ""
        return {"url": url, "title": title, "content": text}
    except (httpx.HTTPError, ValueError, RuntimeError):
        return None


def browse_webpage(url: str) -> dict[str, str]:
    """Read a compact public webpage body for the agent, with guarded redirects."""
    current_url = url
    headers = {"User-Agent": "Mozilla/5.0 (AI-Voice-Chat local demo)", "Accept": "text/html,text/plain;q=0.9,*/*;q=0.2"}
    try:
        with httpx.Client(timeout=20, follow_redirects=False, headers=headers) as client:
            for _ in range(5):
                current_url = validate_public_web_url(current_url)
                with client.stream("GET", current_url) as response:
                    if response.status_code in {301, 302, 303, 307, 308}:
                        location = response.headers.get("location")
                        if not location:
                            raise RuntimeError("网页重定向缺少目标地址.")
                        current_url = urljoin(current_url, location)
                        continue
                    response.raise_for_status()
                    content_type = response.headers.get("content-type", "").lower()
                    if not any(kind in content_type for kind in ("text/", "application/xhtml+xml", "application/json")):
                        raise RuntimeError("该链接不是可读取的文本网页.")
                    raw = bytearray()
                    for chunk in response.iter_bytes():
                        raw.extend(chunk)
                        if len(raw) >= 1_000_000:
                            break
                html = bytes(raw).decode("utf-8", errors="replace")
                title_match = re.search(r"<title[^>]*>([\s\S]*?)</title>", html, flags=re.IGNORECASE)
                title = strip_html(title_match.group(1))[:240] if title_match else ""
                text = strip_html(html)[:12_000]
                if not text:
                    raise RuntimeError("网页没有可读取的正文内容.")
                return {"url": current_url, "title": title, "content": text}
    except httpx.HTTPStatusError as error:
        if error.response.status_code in {403, 429, 451, 468}:
            if fallback := browse_mediawiki_api(str(error.request.url)):
                print(f"[Tool][Browse] HTML blocked; used MediaWiki API for {fallback['url'][:140]!r}", flush=True)
                return fallback
            raise RuntimeError(
                f"该网站拒绝自动读取（HTTP {error.response.status_code}）。智能体应改用其他公开来源或搜索摘要。"
            ) from error
        raise RuntimeError(f"网页读取失败:HTTP {error.response.status_code}") from error
    except httpx.HTTPError as error:
        raise RuntimeError(f"网页读取失败:{error}") from error
    raise RuntimeError("网页重定向次数过多.")


def local_distance_meters(latitude: float, longitude: float, target_latitude: float, target_longitude: float) -> int:
    radius = 6_371_000
    lat1, lon1, lat2, lon2 = map(math.radians, (latitude, longitude, target_latitude, target_longitude))
    value = math.sin((lat2 - lat1) / 2) ** 2 + math.cos(lat1) * math.cos(lat2) * math.sin((lon2 - lon1) / 2) ** 2
    return round(radius * 2 * math.asin(math.sqrt(value)))


def get_local_environment(location: LocationContext | None, query: str, radius_m: int = 1500) -> dict:
    """Fetch current weather and nearby map facilities after explicit browser consent."""
    if location is None:
        raise RuntimeError("未获得用户的位置授权，请先请用户点击“位置”按钮授权。")
    radius_m = max(300, min(int(radius_m), 5000))
    latitude, longitude = location.latitude, location.longitude
    weather_url = (
        "https://api.open-meteo.com/v1/forecast?latitude="
        f"{latitude:.5f}&longitude={longitude:.5f}&current=temperature_2m,apparent_temperature,"
        "relative_humidity_2m,precipitation,weather_code,wind_speed_10m&timezone=auto"
    )
    overpass_query = (
        "[out:json][timeout:15];("
        f"nwr(around:{radius_m},{latitude:.5f},{longitude:.5f})[amenity~\"restaurant|cafe|fast_food|hospital|clinic|pharmacy|bank|atm|toilets|parking|fuel\"];"
        f"nwr(around:{radius_m},{latitude:.5f},{longitude:.5f})[shop~\"supermarket|convenience\"];"
        f"nwr(around:{radius_m},{latitude:.5f},{longitude:.5f})[highway=bus_stop];"
        ");out center 40;"
    )
    weather: dict = {}
    facilities: list[dict[str, object]] = []
    with httpx.Client(timeout=20, headers={"User-Agent": "AI-Voice-Chat local demo"}) as client:
        try:
            weather_response = client.get(weather_url)
            weather_response.raise_for_status()
            weather = weather_response.json().get("current", {})
        except (httpx.HTTPError, ValueError) as error:
            weather = {"error": f"天气查询失败:{error}"}
        try:
            facilities_response = client.post("https://overpass-api.de/api/interpreter", data=overpass_query)
            facilities_response.raise_for_status()
            elements = facilities_response.json().get("elements", [])
            for item in elements if isinstance(elements, list) else []:
                if not isinstance(item, dict):
                    continue
                tags = item.get("tags") if isinstance(item.get("tags"), dict) else {}
                target_latitude = item.get("lat", item.get("center", {}).get("lat"))
                target_longitude = item.get("lon", item.get("center", {}).get("lon"))
                if not isinstance(target_latitude, (int, float)) or not isinstance(target_longitude, (int, float)):
                    continue
                kind = tags.get("amenity") or tags.get("shop") or ("公交站" if tags.get("highway") == "bus_stop" else "设施")
                facilities.append({
                    "name": tags.get("name") or str(kind),
                    "type": kind,
                    "distance_m": local_distance_meters(latitude, longitude, target_latitude, target_longitude),
                })
            facilities.sort(key=lambda item: int(item["distance_m"]))
            facilities = facilities[:20]
        except (httpx.HTTPError, ValueError) as error:
            facilities = [{"error": f"附近设施查询失败:{error}"}]
    return {
        "query": query[:300], "radius_m": radius_m, "weather": weather,
        "facilities": facilities, "accuracy_m": round(location.accuracy),
    }


def compose_user_content(request: ChatRequest, attachments: list[dict], search_results: list[dict[str, str]]) -> str | list[dict]:
    sections = [request.message.strip()]
    if request.location:
        sections.append(
            "【用户已授权本轮定位，仅用于天气与附近设施查询】"
            f"纬度:{request.location.latitude:.5f}，经度:{request.location.longitude:.5f}，"
            f"定位误差约:{round(request.location.accuracy)}米。"
        )
    document_parts = []
    for attachment in attachments:
        if attachment["kind"] != "image":
            document_parts.append(f"【用户上传文件:{attachment['name']}】\n{attachment['text']}")
    if document_parts:
        sections.append("\n\n".join(document_parts))
    if search_results:
        sources = []
        for index, result in enumerate(search_results, start=1):
            source = f"{index}. {result['title']}\n{result['snippet']}\n来源:{result['url']}"
            sources.append(source)
        sections.append("【联网搜索结果,仅作参考;回答中请说明信息来自联网搜索,并按需标注来源】\n" + "\n\n".join(sources))
    text = "\n\n".join(section for section in sections if section)
    image_blocks = [
        {"type": "image_url", "image_url": {"url": f"data:{attachment['content_type']};base64,{attachment['data']}"}}
        for attachment in attachments if attachment["kind"] == "image"
    ]
    if not image_blocks:
        return text
    return [{"type": "text", "text": text}, *image_blocks]


def chat_messages(request: ChatRequest, attachments: list[dict] | None = None, search_results: list[dict[str, str]] | None = None) -> list[BaseMessage]:
    """将网页历史转换成 LangChain 的标准消息对象."""
    system_content = f"{SYSTEM_PROMPT}\n\n{TOOL_SYSTEM_PROMPT}" if request.agent_enabled else SYSTEM_PROMPT
    if live2d_prompt := live2d_control_system_prompt(request.live2d_model_id):
        system_content = f"{system_content}\n\n{live2d_prompt}"
    messages: list[BaseMessage] = [SystemMessage(content=system_content)]
    for item in request.history[-12:]:
        role, content = item.get("role"), item.get("content")
        if role in {"user", "assistant"} and isinstance(content, str) and content.strip():
            if role == "user":
                messages.append(HumanMessage(content=content[:4000]))
            else:
                # The UI intentionally stores only display text.  Rebuild every
                # assistant row into the output protocol before showing it to the
                # model, so a later turn cannot learn a vector-less format.
                messages.append(AIMessage(content=assistant_history_for_model(content)))
    messages.append(HumanMessage(content=compose_user_content(request, attachments or [], search_results or [])))
    return messages


def chunk_text(content: object) -> str:
    """兼容 LangChain 中字符串及内容块两种流式返回格式."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        texts: list[str] = []
        for block in content:
            if isinstance(block, str):
                texts.append(block)
            elif isinstance(block, dict) and isinstance(block.get("text"), str):
                texts.append(block["text"])
        return "".join(texts)
    return ""


def trailing_tag_prefix(text: str, tag: str) -> int:
    """返回末尾与标签开头重叠的长度,处理标签被流式分片截断的情况."""
    for length in range(min(len(text), len(tag) - 1), 0, -1):
        if text.endswith(tag[:length]):
            return length
    return 0


def response_stream_events(chunks: Iterator[str]) -> Iterator[tuple[str, str]]:
    """Split a tagged model stream into visible answer and collapsible thinking events."""
    buffer = ""
    is_thinking = False
    for chunk in chunks:
        buffer += chunk
        while buffer:
            if is_thinking:
                end = buffer.find("</think>")
                if end == -1:
                    keep = trailing_tag_prefix(buffer, "</think>")
                    if len(buffer) > keep:
                        yield "thinking", buffer[:-keep] if keep else buffer
                    buffer = buffer[-keep:] if keep else ""
                    break
                if thought := buffer[:end]:
                    yield "thinking", thought
                buffer = buffer[end + len("</think>"):]
                is_thinking = False
                continue

            start = buffer.find("<think>")
            if start != -1:
                if visible := buffer[:start]:
                    yield "text", visible
                buffer = buffer[start + len("<think>"):]
                is_thinking = True
                continue

            keep = trailing_tag_prefix(buffer, "<think>")
            if len(buffer) > keep:
                yield "text", buffer[:-keep] if keep else buffer
            buffer = buffer[-keep:] if keep else ""

    if buffer:
        yield "thinking" if is_thinking else "text", buffer


def is_internal_agent_payload(value: str) -> bool:
    """Keep model tool traces and accidental JSON out of visible chat and TTS."""
    stripped = value.strip()
    if not stripped:
        return True
    if stripped.startswith("```") or stripped in {"{", "}", "[", "]"}:
        return True
    try:
        parsed = json.loads(stripped)
        if isinstance(parsed, (dict, list)):
            return True
    except json.JSONDecodeError:
        pass
    return bool(re.match(
        r'^["\']?(?:tool_calls?|function_call|function|arguments?|call_id|tool_call_id|name|query|url|image_url|input|output|event|data)["\']?\s*[:=]',
        stripped,
        flags=re.IGNORECASE,
    ))


def split_internal_json_prefix(value: str) -> tuple[object | None, str]:
    """Split a JSON tool payload accidentally followed by a normal reply.

    Some OpenAI-compatible gateways occasionally put a tool result and the
    following natural-language answer in the same content stream.  `json.loads`
    cannot parse that mixed string, so use `raw_decode` to remove only the JSON
    prefix and preserve any actual answer after it.
    """
    stripped = value.strip()
    if not stripped.startswith(("{", "[")):
        return None, value
    try:
        payload, index = json.JSONDecoder().raw_decode(stripped)
    except json.JSONDecodeError:
        return None, value
    if not isinstance(payload, (dict, list)):
        return None, value
    return payload, stripped[index:].lstrip()


def hide_internal_agent_content(value: str) -> tuple[str, str | None]:
    """Return user-visible text plus a small collapsible internal-progress note."""
    payload, visible_tail = split_internal_json_prefix(value)
    if payload is not None:
        return visible_tail, "智能体已接收结构化工具数据，正在整理为可读回答。\n"
    if is_internal_agent_payload(value):
        return "", "智能体正在处理内部工具数据。\n"
    return value, None


def tool_result_thinking_note(event_type: str, payload: dict) -> str:
    """Expose tool progress without leaking the raw JSON result into the reply."""
    if event_type == "local_environment":
        return "智能体已取得当前位置的天气和周边设施数据，正在整理回答。\n"
    if event_type == "webpage":
        title = str(payload.get("title", "")).strip()
        return f"智能体已读取网页{f'《{title}》' if title else ''}，正在提炼重点。\n"
    return "智能体已收到工具结果，正在整理回答。\n"


def parse_live2d_control_suffix(raw_line: str) -> tuple[str, dict | None]:
    """Separate optional `##{...}` metadata before parsing the emotion vector."""
    prefix, delimiter, payload = raw_line.rpartition(LIVE2D_CONTROL_DELIMITER)
    if not delimiter or not payload.lstrip().startswith("{"):
        return raw_line, None
    try:
        parsed = json.loads(payload.strip())
    except json.JSONDecodeError:
        return raw_line, None
    return (prefix, parsed) if isinstance(parsed, dict) else (raw_line, None)


def serialize_emotion_line(text: str, vector: list[float], live2d_control: dict[str, str | None] | None = None) -> str:
    line = f"{text}{EMOTION_DELIMITER}{json.dumps(vector, ensure_ascii=False, separators=(',', ':'))}"
    if live2d_control is not None:
        line += f"{LIVE2D_CONTROL_DELIMITER}{json.dumps(live2d_control, ensure_ascii=False, separators=(',', ':'))}"
    return line


def parse_emotion_line(raw_line: str) -> tuple[str, list[float], dict | None] | None:
    """Read visible text, an emotion vector, and optional Live2D control metadata."""
    raw_line = raw_line.strip()
    if not raw_line:
        return None

    protocol_line, live2d_control = parse_live2d_control_suffix(raw_line)
    text, delimiter, vector_text = protocol_line.rpartition(EMOTION_DELIMITER)
    if not delimiter:
        # Some models occasionally omit one ampersand.  Accept this legacy form
        # while always emitting the canonical && form back into conversation memory.
        text, delimiter, vector_text = protocol_line.rpartition("&")
    if not delimiter or not text.strip():
        if is_internal_agent_payload(protocol_line):
            return None
        return protocol_line, DEFAULT_EMOTION_VECTOR.copy(), live2d_control
    if is_internal_agent_payload(text):
        return None
    try:
        values = json.loads(vector_text.strip().replace(",", ","))
        if not isinstance(values, list) or len(values) != 8:
            raise ValueError("expected 8 values")
        vector = [float(value) for value in values]
        if not all(math.isfinite(value) for value in vector):
            raise ValueError("vector contains a non-finite number")
        return text.strip(), [max(0.0, min(1.0, value)) for value in vector], live2d_control
    except (TypeError, ValueError, json.JSONDecodeError):
        # Keep the answer readable even if a model misses the output protocol.
        return text.strip() or protocol_line, DEFAULT_EMOTION_VECTOR.copy(), live2d_control


def assistant_history_for_model(content: str) -> str:
    """Convert UI-visible assistant text into canonical emotion-tagged history."""
    protocol_lines: list[str] = []
    for raw_line in content.splitlines():
        parsed = parse_emotion_line(raw_line)
        if not parsed:
            continue
        text, vector, live2d_control = parsed
        protocol_lines.append(serialize_emotion_line(text, vector, live2d_control))

    if not protocol_lines:
        return ""

    # Do not cut a final emotion vector in half when applying the history limit.
    kept_lines: list[str] = []
    used = 0
    for line in protocol_lines:
        extra = len(line) + (1 if kept_lines else 0)
        if used + extra > 4000:
            break
        kept_lines.append(line)
        used += extra
    return "\n".join(kept_lines) or protocol_lines[0]


def tool_call_arguments(tool_call: dict) -> dict:
    arguments = tool_call.get("args", {})
    if isinstance(arguments, dict):
        return arguments
    if isinstance(arguments, str):
        try:
            parsed = json.loads(arguments)
            return parsed if isinstance(parsed, dict) else {}
        except json.JSONDecodeError:
            return {}
    return {}


def model_tool_calls(message: object) -> list[dict]:
    """Read both complete LangChain tool calls and streamed tool-call chunks."""
    complete = getattr(message, "tool_calls", []) or []
    calls = [dict(item) for item in complete if isinstance(item, dict)]
    if calls:
        return calls

    grouped: dict[int, dict] = {}
    for chunk in getattr(message, "tool_call_chunks", []) or []:
        if not isinstance(chunk, dict):
            continue
        index = int(chunk.get("index", 0) or 0)
        call = grouped.setdefault(index, {"name": "", "id": "", "args": ""})
        if chunk.get("name"):
            call["name"] += str(chunk["name"])
        if chunk.get("id"):
            call["id"] = str(chunk["id"])
        arguments = chunk.get("args", "")
        if isinstance(arguments, dict):
            call["args"] = arguments
        elif isinstance(call["args"], str):
            call["args"] += str(arguments)
    return [call for _, call in sorted(grouped.items()) if call.get("name")]


def execute_automation_tool(
    tool_call: dict,
    image_config_factory=None,
    reference_images: list[dict] | None = None,
    location: LocationContext | None = None,
) -> tuple[str, dict, str]:
    """Execute one model-selected tool and return its UI event plus ToolMessage content."""
    name = str(tool_call.get("name", ""))
    arguments = tool_call_arguments(tool_call)
    if name == "web_search":
        query = str(arguments.get("query", "")).strip()[:1000]
        if not query:
            raise RuntimeError("联网搜索工具没有收到有效关键词.")
        results = web_search(query)
        payload = {"results": results}
        print(f"[Tool][Search] query={query[:80]!r} results={len(results)}", flush=True)
        return "search_results", payload, json.dumps(payload, ensure_ascii=False)
    if name == "browse_webpage":
        url = str(arguments.get("url", "")).strip()[:2000]
        if not url:
            raise RuntimeError("网页浏览工具没有收到有效链接.")
        payload = browse_webpage(url)
        print(f"[Tool][Browse] url={payload['url'][:140]!r} chars={len(payload['content'])}", flush=True)
        return "webpage", payload, json.dumps(payload, ensure_ascii=False)
    if name == "get_local_environment":
        query = str(arguments.get("query", "")).strip()[:300]
        if not query:
            raise RuntimeError("本地环境工具没有收到查询目标.")
        try:
            radius_m = int(arguments.get("radius_m", 1500))
        except (TypeError, ValueError):
            radius_m = 1500
        payload = get_local_environment(location, query, radius_m)
        print(f"[Tool][Local] query={query[:80]!r} radius={payload['radius_m']}m facilities={len(payload['facilities'])}", flush=True)
        return "local_environment", payload, json.dumps(payload, ensure_ascii=False)
    if name == "generate_image":
        prompt = str(arguments.get("prompt", "")).strip()[:4000]
        if not prompt:
            raise RuntimeError("画图工具没有收到有效提示词.")
        config = image_config_factory() if image_config_factory else resolve_image_config()
        result = generate_image(prompt, config, reference_images)
        print(f"[Tool][Image] generated chars={len(prompt)} references={len(reference_images or [])}", flush=True)
        return "image", result, json.dumps(result, ensure_ascii=False)
    raise RuntimeError(f"不支持的工具调用:{name or 'unknown'}")


def stream_model(
    messages: list[BaseMessage],
    chat_config: dict,
    agent_enabled: bool = True,
    image_config_factory=None,
    reference_images: list[dict] | None = None,
    location: LocationContext | None = None,
) -> Iterator[tuple[str, str, object]]:
    """Stream answers while allowing the model to call search and image tools first."""
    provider = chat_config["provider"]
    api_key = chat_config["api_key"]
    if provider == "lm_studio":
        # LM Studio 通常不验证 API Key,但 LangChain 的 ChatOpenAI 需要非空值.
        api_key = api_key or "lm-studio"
    elif not api_key:
        raise RuntimeError("未配置聊天 API 密钥.请在 config.txt 的 [chat] 中填写 api_key.")
    # Local/LAN OpenAI-compatible gateways (Sub2API included) must bypass a
    # machine-wide SOCKS/HTTP proxy just like LM Studio does.
    use_env_proxy = image_api_uses_env_proxy(chat_config["base_url"])
    http_client = httpx.Client(timeout=300, trust_env=use_env_proxy)
    try:
        model_options = {}
        if chat_config["thinking"]:
            model_options["extra_body"] = (
                {"chat_template_kwargs": {"enable_thinking": True}}
                if provider == "lm_studio" else {"enable_thinking": True}
            )
        print(
            f"[Chat] provider={provider} model={chat_config['model']} "
            f"thinking={chat_config['thinking']} env_proxy={use_env_proxy}",
            flush=True,
        )
        chat_model = ChatOpenAI(
            model=chat_config["model"],
            base_url=chat_config["base_url"].rstrip("/"),
            api_key=api_key,
            temperature=0.7,
            timeout=300,
            max_retries=2,
            stream_usage=False,
            http_client=http_client,
            **model_options,
        )
        full_thinking: list[str] = []
        full_raw_text: list[str] = []
        working_messages = list(messages)
        tool_model = chat_model.bind_tools(AUTOMATION_TOOLS) if agent_enabled else chat_model
        max_rounds = MAX_AGENT_TOOL_ROUNDS if agent_enabled else 1
        tool_counts: dict[str, int] = {}

        for tool_round in range(max_rounds):
            final_chunk = None

            def raw_chunks() -> Iterator[str]:
                nonlocal final_chunk
                for chunk in tool_model.stream(working_messages):
                    final_chunk = chunk if final_chunk is None else final_chunk + chunk
                    if content := chunk_text(chunk.content):
                        yield content

            text_buffer = ""
            for event_type, delta in response_stream_events(raw_chunks()):
                if event_type == "thinking":
                    full_thinking.append(delta)
                    yield "thinking", delta, None
                    continue

                full_raw_text.append(delta)
                text_buffer += delta
                while "\n" in text_buffer:
                    raw_line, text_buffer = text_buffer.split("\n", 1)
                    visible_line, internal_note = hide_internal_agent_content(raw_line)
                    if internal_note:
                        full_thinking.append(internal_note)
                        yield "thinking", internal_note, None
                    parsed = parse_emotion_line(visible_line)
                    if parsed:
                        text, vector, live2d_control = parsed
                        yield "text", text, {"emotion_vector": vector, "live2d_control": live2d_control}
            visible_line, internal_note = hide_internal_agent_content(text_buffer)
            if internal_note:
                full_thinking.append(internal_note)
                yield "thinking", internal_note, None
            parsed = parse_emotion_line(visible_line)
            if parsed:
                text, vector, live2d_control = parsed
                yield "text", text, {"emotion_vector": vector, "live2d_control": live2d_control}

            tool_calls = model_tool_calls(final_chunk)
            if not agent_enabled or not tool_calls:
                break
            if tool_round == max_rounds - 1:
                raise RuntimeError("模型连续请求工具超过上限,已停止本轮调用.")

            normalized_calls: list[dict] = []
            for raw_call in tool_calls:
                if not isinstance(raw_call, dict):
                    continue
                call = dict(raw_call)
                call["id"] = str(call.get("id") or uuid.uuid4().hex)
                call["args"] = tool_call_arguments(call)
                normalized_calls.append(call)
            if not normalized_calls:
                break

            assistant_content = chunk_text(getattr(final_chunk, "content", "")) if final_chunk else ""
            working_messages.append(AIMessage(content=assistant_content, tool_calls=normalized_calls))
            for tool_call in normalized_calls:
                tool_name = str(tool_call.get("name", ""))
                arguments = tool_call_arguments(tool_call)
                tool_counts[tool_name] = tool_counts.get(tool_name, 0) + 1
                limit = {
                    "generate_image": MAX_AGENT_IMAGE_CALLS,
                    "web_search": MAX_AGENT_SEARCH_CALLS,
                    "browse_webpage": MAX_AGENT_BROWSE_CALLS,
                    "get_local_environment": MAX_AGENT_LOCATION_CALLS,
                }.get(tool_name, 0)
                detail = str(
                    arguments.get("query") if tool_name in {"web_search", "get_local_environment"}
                    else arguments.get("url") if tool_name == "browse_webpage"
                    else arguments.get("prompt", "")
                ).strip()[:120]
                step_payload = {"call_id": tool_call["id"], "tool": tool_name, "detail": detail}
                if limit and tool_counts[tool_name] > limit:
                    message = f"本轮智能体最多可调用 {limit} 次该工具."
                    yield "agent_step", json.dumps({**step_payload, "state": "limited", "message": message}, ensure_ascii=False), None
                    working_messages.append(ToolMessage(content=json.dumps({"error": message}, ensure_ascii=False), tool_call_id=tool_call["id"]))
                    continue
                yield "agent_step", json.dumps({**step_payload, "state": "running"}, ensure_ascii=False), None
                status = {
                    "web_search": "模型正在联网搜索…",
                    "browse_webpage": "模型正在阅读网页正文…",
                    "get_local_environment": "模型正在查询本地天气和设施…",
                }.get(tool_name, "模型正在生成图片…")
                yield "tool_status", status, None
                try:
                    event_name, event_payload, tool_result = execute_automation_tool(
                        tool_call, image_config_factory, reference_images, location,
                    )
                    yield event_name, json.dumps(event_payload, ensure_ascii=False), None
                    yield "agent_step", json.dumps({**step_payload, "state": "completed"}, ensure_ascii=False), None
                except Exception as error:
                    message = str(error)
                    tool_result = json.dumps({"error": message}, ensure_ascii=False)
                    yield "agent_step", json.dumps({**step_payload, "state": "failed", "message": message}, ensure_ascii=False), None
                    if tool_name == "web_search":
                        yield "search_results", json.dumps({"results": [], "error": message}, ensure_ascii=False), None
                    elif tool_name == "browse_webpage":
                        yield "webpage", json.dumps({"error": message}, ensure_ascii=False), None
                    elif tool_name == "get_local_environment":
                        yield "local_environment", json.dumps({"error": message}, ensure_ascii=False), None
                    elif tool_name == "generate_image":
                        yield "image", json.dumps({"error": message}, ensure_ascii=False), None
                working_messages.append(ToolMessage(content=tool_result, tool_call_id=tool_call["id"]))
        print("[LLM] stream completed", flush=True)
        if full_thinking:
            print(f"[LLM][thinking]\n{''.join(full_thinking)}", flush=True)
        print(f"[LLM][text]\n{''.join(full_raw_text)}", flush=True)
    except Exception as error:
        raise RuntimeError(f"LangChain 聊天模型调用失败:{str(error)[:500]}") from error
    finally:
        http_client.close()


def tts_text_char_count(text: str) -> int:
    """Count speech content without line-break formatting."""
    return len(re.sub(r"\s+", "", text))


def average_tts_emotion_vector(segments: list[tuple[int, str, list[float]]]) -> list[float]:
    """Average merged-line emotions using each line's spoken-character weight."""
    if not segments:
        return DEFAULT_EMOTION_VECTOR.copy()

    totals = [0.0] * len(DEFAULT_EMOTION_VECTOR)
    total_weight = 0
    for _, text, vector in segments:
        weight = max(1, tts_text_char_count(text))
        safe_vector = vector if len(vector) == len(DEFAULT_EMOTION_VECTOR) else DEFAULT_EMOTION_VECTOR
        for index, value in enumerate(safe_vector):
            totals[index] += float(value) * weight
        total_weight += weight

    return [max(0.0, min(1.0, value / total_weight)) for value in totals]


@app.get("/")
def home():
    return FileResponse(STATIC_DIR / "index.html")


@app.post("/api/attachments")
async def upload_attachment(file: UploadFile = File(...)):
    """Keep the current-turn attachment in process memory; nothing is written to chat history."""
    filename = Path(file.filename or "未命名文件").name.strip()[:160] or "未命名文件"
    kind = attachment_kind(filename, file.content_type)
    try:
        data = await file.read(MAX_ATTACHMENT_BYTES + 1)
    finally:
        await file.close()
    if not data:
        raise HTTPException(status_code=400, detail="不能上传空文件.")
    if len(data) > MAX_ATTACHMENT_BYTES:
        raise HTTPException(status_code=413, detail="单个附件不能超过 10 MB.")

    content_type = file.content_type or "application/octet-stream"
    item = {
        "id": uuid.uuid4().hex,
        "name": filename,
        "kind": kind,
        "size": len(data),
        "created_at": time.time(),
        "content_type": content_type,
    }
    if kind == "image":
        mime_by_suffix = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".webp": "image/webp", ".gif": "image/gif"}
        item["content_type"] = content_type if content_type.startswith("image/") else mime_by_suffix[Path(filename).suffix.lower()]
        item["data"] = base64.b64encode(data).decode("ascii")
    else:
        item["text"] = extract_attachment_text(kind, data)

    clean_expired_attachments()
    with ATTACHMENTS_LOCK:
        ATTACHMENTS[item["id"]] = item
    print(f"[Attachment] received kind={kind} name={filename!r} bytes={len(data)}", flush=True)
    response = {key: item[key] for key in ("id", "name", "kind", "size")}
    if kind == "image":
        response["preview_url"] = f"/api/attachments/{item['id']}/preview"
    return response


@app.get("/api/attachments/{attachment_id}/preview")
def attachment_preview(attachment_id: str):
    clean_expired_attachments()
    if not re.fullmatch(r"[0-9a-f]{32}", attachment_id):
        raise HTTPException(status_code=404, detail="Attachment not found.")
    with ATTACHMENTS_LOCK:
        item = ATTACHMENTS.get(attachment_id)
    if not item or item.get("kind") != "image":
        raise HTTPException(status_code=404, detail="Image attachment not found.")
    return Response(
        content=base64.b64decode(item["data"]),
        media_type=item["content_type"],
        headers={"Cache-Control": "no-store"},
    )


@app.post("/api/images")
def create_image(request: ImageRequest):
    chat_config = resolve_chat_config_from_web(request.web_chat)
    config = resolve_image_config(request.web_image, chat_config)
    prompt = request.prompt.strip()
    print(f"[Image] generating model={config['model']} chars={len(prompt)}", flush=True)
    result = generate_image(prompt, config)
    print("[Image] generation completed", flush=True)
    return result


@app.get("/api/status")
def status():
    voices = TTS.voices()
    chat_config = SETTINGS["chat"]
    image_config = SETTINGS["image"]
    return {
        "configured": bool(chat_config["base_url"] and chat_config["model"] and (chat_config["api_key"] or chat_config["provider"] == "lm_studio")),
        "provider": chat_config["provider"],
        "model": chat_config["model"],
        "chat_defaults": {
            "provider": chat_config["provider"],
            "base_url": chat_config["base_url"],
            "model": chat_config["model"],
            "api_key_configured": bool(chat_config["api_key"]),
            "thinking": chat_config["thinking"],
        },
        "image_defaults": {
            "base_url": image_config["base_url"],
            "model": image_config["model"],
            "api_mode": image_config["api_mode"],
            "responses_model": image_config["responses_model"],
            "api_key_configured": bool(image_config["api_key"]),
        },
        "voices": voices,
        "default_voice": SETTINGS.get("default_voice", "") if SETTINGS.get("default_voice") in voices else (voices[0] if voices else ""),
        "index_tts_home": str(TTS.project_home),
        "playback_prebuffer_segments": SETTINGS["playback_prebuffer_segments"],
        "tts_model_state": TTS.model_status(),
    }


@app.get("/api/live2d/models")
def list_live2d_models():
    models = local_live2d_models()
    return {"models": models, "default_model": models[0] if models else None}


@app.get("/api/conversations")
def list_conversations():
    with CONVERSATIONS_LOCK:
        records = load_saved_conversations()
    records.sort(key=lambda record: float(record.get("updated_at", 0)), reverse=True)
    return {"conversations": [conversation_summary(record) for record in records]}


@app.get("/api/conversations/{conversation_id}")
def get_conversation(conversation_id: str):
    conversation_id = validate_conversation_id(conversation_id)
    with CONVERSATIONS_LOCK:
        records = load_saved_conversations()
    for record in records:
        if record["id"] == conversation_id:
            return record
    raise HTTPException(status_code=404, detail="Conversation not found.")


@app.delete("/api/conversations/{conversation_id}")
def delete_conversation(conversation_id: str):
    conversation_id = validate_conversation_id(conversation_id)
    with CONVERSATIONS_LOCK:
        records = load_saved_conversations()
        remaining = [record for record in records if record["id"] != conversation_id]
        if len(remaining) == len(records):
            raise HTTPException(status_code=404, detail="Conversation not found.")
        write_saved_conversations(remaining)
    return {"id": conversation_id}


@app.post("/api/conversations")
def save_conversation(request: ConversationSaveRequest):
    conversation_id = validate_conversation_id(request.id)
    messages: list[dict] = []
    for message in request.messages:
        if message.role not in {"user", "assistant"}:
            raise HTTPException(status_code=400, detail="Invalid conversation role.")
        messages.append(message.model_dump(exclude_none=True))

    now = time.time()
    with CONVERSATIONS_LOCK:
        records = load_saved_conversations()
        existing = next((record for record in records if record["id"] == conversation_id), None)
        record = {
            "id": conversation_id,
            "title": request.title.strip(),
            "created_at": existing.get("created_at", now) if existing else now,
            "updated_at": now,
            "messages": messages,
        }
        records = [item for item in records if item["id"] != conversation_id]
        records.insert(0, record)
        write_saved_conversations(records[:80])
    return conversation_summary(record)


@app.post("/api/chat")
def chat(request: ChatRequest):
    if request.voice not in TTS.voices():
        raise HTTPException(status_code=400, detail="音色不可用.")

    def generate():
        answer = ""
        model_history = ""
        line_number = 0
        cancel_event = threading.Event()
        completed = False
        model_events: queue.Queue[tuple[str, str, object]] = queue.Queue()
        tts_jobs: queue.Queue[tuple[str, list[float], list[tuple[int, str]], float] | None] = queue.Queue()
        tts_events: queue.Queue[tuple[str, dict]] = queue.Queue()
        attachments: list[dict] = []
        search_results: list[dict[str, str]] = []
        pending_tts_segments: list[tuple[int, str, list[float]]] = []
        pending_tts_started_at: float | None = None

        try:
            if request.attachments:
                yield sse("status", {"message": "正在读取上传的文件…"})
                attachments = resolve_attachments(request.attachments)
                yield sse("status", {"message": f"已读取 {len(attachments)} 个附件,正在准备模型请求…"})
            if request.web_search:
                yield sse("status", {"message": "正在联网搜索…"})
                try:
                    search_results = web_search(request.message)
                    yield sse("search_results", {"results": search_results})
                    yield sse("status", {"message": f"联网搜索完成,找到 {len(search_results)} 条参考结果."})
                    print(f"[Search] query={request.message[:80]!r} results={len(search_results)}", flush=True)
                except Exception as error:
                    message = str(error)
                    yield sse("search_results", {"results": [], "error": message})
                    yield sse("status", {"message": "联网搜索暂不可用,将直接询问模型."})
                    print(f"[Search] failed: {message}", flush=True)
        except Exception as error:
            yield sse("error", {"message": str(error)})
            return

        def model_worker() -> None:
            try:
                messages = chat_messages(request, attachments, search_results)
                chat_config = resolve_chat_config(request)
                image_config_factory = lambda: resolve_image_config(request.web_image, chat_config)
                reference_images = [
                    {"content_type": item["content_type"], "data": item["data"]}
                    for item in attachments if item["kind"] == "image"
                ]
                for event_type, delta, metadata in stream_model(
                    messages, chat_config, request.agent_enabled, image_config_factory, reference_images, request.location,
                ):
                    if cancel_event.is_set():
                        return
                    model_events.put((event_type, delta, metadata))
            except Exception as error:
                model_events.put(("model_error", str(error), None))
            finally:
                model_events.put(("model_done", "", None))

        def tts_worker() -> None:
            completed_segments = 0
            while True:
                job = tts_jobs.get()
                try:
                    if job is None:
                        return
                    segment, emotion_vector, line_segments, speaking_speed = job
                    if cancel_event.is_set():
                        continue
                    line_ids = [line_id for line_id, _ in line_segments]
                    current_line_id = line_ids[-1] if line_ids else 1
                    batch_number = completed_segments + 1
                    model_state = TTS.model_status()
                    tts_events.put(("tts_progress", {
                        "stage": "loading" if model_state in {"not_loaded", "loading"} else "generating",
                        "line_id": current_line_id,
                        "line_ids": line_ids,
                        "batch_number": batch_number,
                        "completed": completed_segments,
                    }))
                    started_at = time.perf_counter()
                    print(f"[TTS] start lines={line_ids} chars={tts_text_char_count(segment)} text={segment[:72]!r}", flush=True)
                    audio_name = TTS.synthesize(segment, request.voice, emotion_vector, speaking_speed)
                    if cancel_event.is_set():
                        continue
                    elapsed = time.perf_counter() - started_at
                    print(f"[TTS] ready lines={line_ids} chars={tts_text_char_count(segment)} seconds={elapsed:.2f}", flush=True)
                    completed_segments += 1
                    tts_events.put(("audio", {
                        "audio": f"/media/{audio_name}",
                        "text": segment,
                        "line_id": current_line_id,
                        "line_ids": line_ids,
                        "segments": [{"line_id": line_id, "text": text} for line_id, text in line_segments],
                        "emotion_vector": emotion_vector,
                        "speaking_speed": speaking_speed,
                    }))
                    tts_events.put(("tts_progress", {
                        "stage": "ready",
                        "line_id": current_line_id,
                        "line_ids": line_ids,
                        "batch_number": batch_number,
                        "completed": completed_segments,
                    }))
                except Exception as error:
                    tts_events.put(("tts_error", {"message": str(error)}))
                finally:
                    tts_jobs.task_done()

        model_thread = threading.Thread(target=model_worker, name="chat-stream", daemon=True)
        tts_thread = threading.Thread(target=tts_worker, name="indextts-stream", daemon=True)
        model_thread.start()
        tts_thread.start()

        def ready_tts_events() -> Iterator[bytes]:
            while True:
                try:
                    event, data = tts_events.get_nowait()
                except queue.Empty:
                    return
                yield sse(event, data)

        def enqueue_tts_batch(segments: list[tuple[int, str, list[float]]]) -> Iterator[bytes]:
            if not segments or cancel_event.is_set():
                return
            text = "\n".join(segment_text for _, segment_text, _ in segments)
            line_segments = [(line_id, segment_text) for line_id, segment_text, _ in segments]
            emotion_vector = average_tts_emotion_vector(segments)
            tts_jobs.put((text, emotion_vector, line_segments, request.speaking_speed))
            print(
                f"[TTS] queued lines={[line_id for line_id, _ in line_segments]} "
                f"chars={tts_text_char_count(text)} text={text[:72]!r}",
                flush=True,
            )
            yield sse("status", {"message": "正在预生成下一段语音…"})

        def flush_pending_tts() -> Iterator[bytes]:
            nonlocal pending_tts_segments, pending_tts_started_at
            if not pending_tts_segments:
                return
            segments = pending_tts_segments
            pending_tts_segments = []
            pending_tts_started_at = None
            yield from enqueue_tts_batch(segments)

        def accept_tts_line(text: str, emotion_vector: list[float], line_id: int) -> Iterator[bytes]:
            """Merge only consecutive short visible lines; display rows stay intact."""
            nonlocal pending_tts_started_at
            segment = text.strip()
            if not segment or cancel_event.is_set():
                return

            candidate = (line_id, segment, emotion_vector)
            candidate_chars = tts_text_char_count(segment)
            if candidate_chars > TTS_SHORT_LINE_CHARS:
                # A normal-length row is suitable alone. Flush earlier short
                # rows first so both audio and highlights remain in order.
                yield from flush_pending_tts()
                yield from enqueue_tts_batch([candidate])
                return

            pending_chars = sum(tts_text_char_count(item[1]) for item in pending_tts_segments)
            if pending_tts_segments and pending_chars + candidate_chars > TTS_BATCH_MAX_CHARS:
                yield from flush_pending_tts()
                pending_chars = 0

            pending_tts_segments.append(candidate)
            if pending_tts_started_at is None:
                pending_tts_started_at = time.monotonic()
            if pending_chars + candidate_chars >= TTS_BATCH_TARGET_CHARS:
                yield from flush_pending_tts()

        try:
            yield sse("status", {"message": "正在获取回答…"})
            model_finished = False
            tts_closed = False
            while True:
                # Audio events are checked continuously, even while the language
                # model is paused between two stream chunks.
                yield from ready_tts_events()
                if (
                    pending_tts_segments
                    and pending_tts_started_at is not None
                    and time.monotonic() - pending_tts_started_at >= TTS_BATCH_MAX_WAIT_SECONDS
                ):
                    # Keep streaming responsive if the next newline arrives late.
                    yield from flush_pending_tts()
                if model_finished and not tts_closed:
                    answer = answer.strip()
                    yield from flush_pending_tts()
                    tts_jobs.put(None)
                    tts_closed = True

                if model_finished and not tts_thread.is_alive() and tts_events.empty():
                    break

                try:
                    event_type, payload, metadata = model_events.get(timeout=0.12)
                except queue.Empty:
                    continue

                if event_type == "model_done":
                    model_finished = True
                    continue
                if event_type == "model_error":
                    raise RuntimeError(str(payload))

                delta = str(payload)
                if event_type == "tool_status":
                    yield sse("status", {"message": delta})
                    continue
                if event_type in {"search_results", "image", "agent_step", "webpage", "local_environment"}:
                    try:
                        tool_payload = json.loads(delta)
                    except json.JSONDecodeError:
                        tool_payload = {"error": "工具返回数据无法解析."}
                    if event_type in {"webpage", "local_environment"}:
                        # Tool results can be very large JSON objects.  They are
                        # context for the model, not an answer or TTS input.
                        yield sse("thinking", {"text": tool_result_thinking_note(event_type, tool_payload)})
                    else:
                        yield sse(event_type, tool_payload)
                    continue
                if event_type == "thinking":
                    yield sse("thinking", {"text": delta})
                    continue
                line_number += 1
                answer = f"{answer}\n{delta}" if answer else delta
                raw_vector = metadata.get("emotion_vector") if isinstance(metadata, dict) else metadata
                vector = raw_vector if isinstance(raw_vector, list) and len(raw_vector) == 8 else DEFAULT_EMOTION_VECTOR.copy()
                raw_live2d_control = metadata.get("live2d_control") if isinstance(metadata, dict) else None
                live2d_control = normalize_live2d_control(raw_live2d_control, request.live2d_model_id)
                history_line = serialize_emotion_line(delta, vector, live2d_control)
                model_history = f"{model_history}\n{history_line}" if model_history else history_line
                yield sse("delta", {
                    "text": delta,
                    "line_id": line_number,
                    "emotion_vector": vector,
                    "live2d_control": live2d_control,
                })
                yield from accept_tts_line(delta, vector, line_number)
            completed = True
            yield sse("done", {"answer": answer, "model_history": model_history})
        except Exception as error:
            yield sse("error", {"message": str(error)})
        finally:
            cancel_event.set()
            if not completed:
                print("[Chat] client stream closed; cancelling pending model and TTS work.", flush=True)
            if tts_thread.is_alive():
                tts_jobs.put(None)

    return StreamingResponse(generate(), media_type="text/event-stream", headers={"Cache-Control": "no-cache"})
